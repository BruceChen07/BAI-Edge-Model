from __future__ import annotations

import logging
import time
from pathlib import Path

from docx import Document
from openpyxl import Workbook

from app.core.config import config
from app.core.database import get_connection
from app.core.logging import get_logger, log_event
from app.core.request_context import get_trace_id
from app.schemas.export import ExportRequestDTO
from app.utils.ids import new_id


class ExportService:
    def __init__(self) -> None:
        self.logger = get_logger("app.export")

    def create_markdown(self, payload: ExportRequestDTO) -> dict:
        started_at = time.time()
        content = self._resolve_content(payload.source_type, payload.source_id)
        export_id = new_id("exp")
        file_name = f"{payload.title or 'export'}-{export_id}.md"
        target = config.exports_root / file_name
        target.write_text(content, encoding="utf-8")
        self._register(export_id, "markdown", payload, target)
        self._log_export("markdown", payload, export_id, target, started_at)
        return {"export_id": export_id, "file_name": file_name, "file_path": str(target)}

    def create_docx(self, payload: ExportRequestDTO) -> dict:
        started_at = time.time()
        content = self._resolve_content(payload.source_type, payload.source_id)
        export_id = new_id("exp")
        file_name = f"{payload.title or 'export'}-{export_id}.docx"
        target = config.exports_root / file_name
        document = Document()
        document.add_heading(payload.title or "Export", level=1)
        for paragraph in content.split("\n\n"):
            document.add_paragraph(paragraph)
        document.save(target)
        self._register(export_id, "docx", payload, target)
        self._log_export("docx", payload, export_id, target, started_at)
        return {"export_id": export_id, "file_name": file_name, "file_path": str(target)}

    def create_xlsx(self, payload: ExportRequestDTO) -> dict:
        started_at = time.time()
        content = self._resolve_content(payload.source_type, payload.source_id)
        export_id = new_id("exp")
        file_name = f"{payload.title or 'export'}-{export_id}.xlsx"
        target = config.exports_root / file_name
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Export"
        for index, line in enumerate(content.splitlines(), start=1):
            sheet.cell(row=index, column=1, value=line)
        workbook.save(target)
        self._register(export_id, "xlsx", payload, target)
        self._log_export("xlsx", payload, export_id, target, started_at)
        return {"export_id": export_id, "file_name": file_name, "file_path": str(target)}

    def get_export(self, export_id: str) -> dict:
        with get_connection() as conn:
            row = conn.execute(
                "SELECT * FROM exports WHERE id = ?", (export_id,)).fetchone()
        return dict(row)

    def _register(self, export_id: str, export_type: str, payload: ExportRequestDTO, target: Path) -> None:
        with get_connection() as conn:
            conn.execute(
                "INSERT INTO exports (id, export_type, source_type, source_id, file_name, file_path, status) VALUES (?, ?, ?, ?, ?, ?, 'ready')",
                (export_id, export_type, payload.source_type,
                 payload.source_id, target.name, str(target)),
            )

    def _resolve_content(self, source_type: str, source_id: str) -> str:
        with get_connection() as conn:
            if source_type == "session":
                rows = conn.execute(
                    "SELECT role, content FROM messages WHERE session_id = ? ORDER BY created_at ASC", (source_id,)).fetchall()
                return "\n\n".join(f"## {row['role']}\n{row['content']}" for row in rows)
            if source_type == "knowledge_base":
                rows = conn.execute(
                    "SELECT file_name, storage_path FROM documents WHERE kb_id = ? ORDER BY created_at DESC", (source_id,)).fetchall()
                return "\n".join(f"- {row['file_name']} ({row['storage_path']})" for row in rows)
            if source_type == "agent_run":
                rows = conn.execute(
                    "SELECT step_no, tool_name, tool_output FROM agent_steps WHERE run_id = ? ORDER BY step_no ASC", (source_id,)).fetchall()
                return "\n\n".join(
                    f"Step {row['step_no']} - {row['tool_name']}\n{row['tool_output']}"
                    for row in rows
                )
        return f"Unsupported source: {source_type} / {source_id}"

    def _log_export(self, export_type: str, payload: ExportRequestDTO, export_id: str, target: Path, started_at: float) -> None:
        log_event(
            self.logger,
            logging.INFO,
            "export.created",
            "Export file created",
            module_name="export",
            trace_id=get_trace_id(),
            export_type=export_type,
            export_id=export_id,
            source_type=payload.source_type,
            source_id=payload.source_id,
            file_path=str(target),
            elapsed_ms=round((time.time() - started_at) * 1000, 2),
        )
