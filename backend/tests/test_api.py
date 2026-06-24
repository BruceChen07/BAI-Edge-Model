from __future__ import annotations

import io
import logging
from pathlib import Path

import httpx
from fastapi.testclient import TestClient
from docx import Document
from openpyxl import Workbook

from app.core.database import init_db
from app.core.config import config
from app.core.logging import setup_logging
from app.schemas.chat import ChatResponseDTO
from app.main import app, chat_attachment_service, chat_service, ingest_service

test_db_path = config.project_root / "storage" / "test_bai_edge_model.db"
object.__setattr__(config, "db_path", test_db_path)
if config.db_path.exists():
    config.db_path.unlink()

client = TestClient(app)
init_db()
setup_logging()


def _create_docx_bytes() -> bytes:
    buffer = io.BytesIO()
    document = Document()
    document.add_heading("Deployment", level=1)
    document.add_paragraph("Edge deployment with local rag memory.")
    document.save(buffer)
    return buffer.getvalue()


def _create_xlsx_bytes() -> bytes:
    buffer = io.BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Plan"
    sheet["A1"] = "Topic"
    sheet["B1"] = "Value"
    sheet["A2"] = "RAG"
    sheet["B2"] = "Local retrieval"
    workbook.save(buffer)
    return buffer.getvalue()


def _screenshot_like_ocr_text() -> str:
    return """# 1.6 深入剖析GPT架构

在本章之前，我们提到了 GPT 类模型、GPT-3 和 ChatGPT。现在让我们更深入地了解通用的 GPT 架构。
首先，GPT 是“生成预训练变换器”（Generative Pretrained Transformer）的缩写，最初是在以下论文中提出的：

- 《通过生成预训练改善语言理解》（2018）是由 OpenAI 的 Radford 等人撰写的，链接：http://cdn.openai.com/researchcovers/language-unsupervised/language_understanding_paper.pdf
""".strip()


def _tiny_png_bytes() -> bytes:
    return (
        b"\x89PNG\r\n\x1a\n"
        b"\x00\x00\x00\rIHDR"
        b"\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00"
        b"\x90wS\xde"
        b"\x00\x00\x00\x0cIDAT\x08\xd7c\xf8\x0f\x00\x01\x01\x01\x00"
        b"\x18\xdd\x8d\xb4"
        b"\x00\x00\x00\x00IEND\xaeB`\x82"
    )


def test_system_info() -> None:
    response = client.get("/api/v1/system/info")
    assert response.status_code == 200
    assert response.json()["data"]["app_name"] == "BAI Edge Model"


def test_models_endpoint(monkeypatch) -> None:
    monkeypatch.setattr(
        chat_service.ollama_service,
        "list_models",
        lambda: [
            {"name": "qwen3.5:9b", "size": 1, "modified_at": "now", "digest": "abc"},
            {"name": "qwen3.5:4b", "size": 1, "modified_at": "now", "digest": "def"},
        ],
    )
    monkeypatch.setattr(
        "app.main.llmfit_service.get_model_info",
        lambda model_name: None
        if model_name != "qwen3.5:9b"
        else type(
            "Rec",
            (),
            {
                "provider": "Qwen",
                "param_size": "9B",
                "score": type(
                    "Score",
                    (),
                    {
                        "total": 88.0,
                        "quality": 86.0,
                        "speed": 74.0,
                        "fit": 81.0,
                        "context": 79.0,
                    },
                )(),
                "fit_level": "good",
                "estimated_tps": 12.0,
                "quantization": "Q4_K_M",
                "memory_required_gb": 6.6,
                "run_mode": "GPU",
                "use_case": "general",
                "max_context": 32768,
                "is_moe": False,
            },
        )(),
    )
    response = client.get("/api/v1/models")
    assert response.status_code == 200
    assert response.json()["data"][0]["name"] == "qwen3.5:9b"
    assert response.json()["data"][0]["provider"] == "Qwen"
    assert response.json()["data"][0]["fit_level"] == "good"
    assert response.json()["data"][0]["estimated_tps"] == 12.0
    assert response.json()["data"][0]["source"] == "local+llmfit"
    assert response.json()["data"][0]["supports_multimodal"] is False
    assert response.json()["data"][0]["supports_file_upload"] is False
    assert response.json()["data"][0]["supported_upload_types"] == []


def test_models_endpoint_exposes_multimodal_capability(monkeypatch) -> None:
    monkeypatch.setattr(
        chat_service.ollama_service,
        "list_models",
        lambda: [
            {"name": "llava:7b", "size": 1, "modified_at": "now", "digest": "vision"},
        ],
    )
    monkeypatch.setattr(
        "app.main.llmfit_service.get_model_info", lambda _: None)

    response = client.get("/api/v1/models")

    assert response.status_code == 200
    payload = response.json()["data"][0]
    assert payload["name"] == "llava:7b"
    assert payload["supports_multimodal"] is True
    assert payload["supports_file_upload"] is True
    assert payload["supported_upload_types"] == ["document", "image"]
    assert payload["capability_source"] == "name_inference"


def test_catalog_sync_local_upserts_models(monkeypatch) -> None:
    monkeypatch.setattr(
        chat_service.ollama_service,
        "list_models",
        lambda: [
            {"name": "qwen3.5:9b", "size": 1, "modified_at": "now", "digest": "abc"},
            {"name": "llama3.2:3b", "size": 1,
                "modified_at": "now", "digest": "def"},
            {"name": "gemma4:12b", "size": 1, "modified_at": "now", "digest": "ghi"},
        ],
    )
    monkeypatch.setattr(
        "app.main.llmfit_service.get_model_info",
        lambda model_name: None
        if model_name != "qwen3.5:9b"
        else type(
            "Rec",
            (),
            {
                "provider": "Qwen",
                "param_size": "9B",
                "score": type(
                    "Score",
                    (),
                    {
                        "total": 90.0,
                        "quality": 88.0,
                        "speed": 70.0,
                        "fit": 84.0,
                        "context": 80.0,
                    },
                )(),
                "fit_level": "good",
                "estimated_tps": 11.0,
                "quantization": "Q4_K_M",
                "memory_required_gb": 6.6,
                "run_mode": "GPU",
                "use_case": "general",
                "max_context": 32768,
                "is_moe": False,
            },
        )(),
    )

    sync_response = client.post(
        "/api/v1/catalog/sync", json={"source": "local"})

    assert sync_response.status_code == 200
    assert sync_response.json(
    )["data"]["message"] == "local model sync completed"

    list_response = client.get(
        "/api/v1/catalog", params={"limit": 100, "sort_by": "model_name", "sort_dir": "asc"})
    assert list_response.status_code == 200
    payload = list_response.json()["data"]
    items = payload["items"]
    names = [i["model_name"] for i in items]
    assert "qwen3.5:9b" in names
    assert "llama3.2:3b" in names
    assert "gemma4:12b" in names
    synced_item = next(i for i in items if i["model_name"] == "qwen3.5:9b")
    assert synced_item["source"] == "local"
    assert synced_item["available"] is True
    assert synced_item["fit_level"] == "good"
    assert synced_item["estimated_tps"] == 11.0


def test_knowledge_base_chat_export_flow(monkeypatch) -> None:
    monkeypatch.setattr(
        chat_service.ollama_service,
        "list_models",
        lambda: [{"name": "qwen3.5:4b", "size": 1,
                  "modified_at": "now", "digest": "abc"}],
    )
    monkeypatch.setattr(
        chat_service.ollama_service,
        "chat",
        lambda **kwargs: f"mocked ollama reply from {kwargs['model_name']}",
    )

    kb_response = client.post(
        "/api/v1/knowledge-bases",
        json={"name": "Demo KB", "description": "test"},
    )
    assert kb_response.status_code == 200
    kb_id = kb_response.json()["data"]["id"]

    upload_response = client.post(
        f"/api/v1/knowledge-bases/{kb_id}/files/upload?enable_ocr=true",
        files={
            "file": (
                "demo.txt",
                io.BytesIO(b"edge deployment rag agent memory export"),
                "text/plain",
            )
        },
    )
    assert upload_response.status_code == 200

    session_response = client.post(
        "/api/v1/sessions", json={"title": "Demo Session"})
    assert session_response.status_code == 200
    session_id = session_response.json()["data"]["id"]

    chat_response = client.post(
        "/api/v1/chat/completions",
        json={
            "session_id": session_id,
            "query": "deployment rag",
            "model_name": "qwen3.5:4b",
            "knowledge_base_ids": [kb_id],
        },
    )
    assert chat_response.status_code == 200
    assert "mocked ollama reply from qwen3.5:4b" in chat_response.json()[
        "data"]["answer"]
    assert chat_response.json()["data"]["model_used"] == "qwen3.5:4b"
    citations = chat_response.json()["data"]["citations"]
    assert citations
    assert citations[0]["file_name"] == "demo.txt"
    assert "source_label" in citations[0]
    assert citations[0]["chunk_index"] >= 0
    assert "quote_excerpt" in citations[0]
    assert citations[0]["matched_terms"]

    export_response = client.post(
        "/api/v1/exports/markdown",
        json={
            "source_type": "session",
            "source_id": session_id,
            "title": "session-export",
        },
    )
    assert export_response.status_code == 200
    export_path = Path(export_response.json()["data"]["file_path"])
    assert export_path.exists()


def test_chat_completions_no_longer_depends_on_debug_server(monkeypatch) -> None:
    session_response = client.post(
        "/api/v1/sessions", json={"title": "Debug Hook Regression"})
    assert session_response.status_code == 200
    session_id = session_response.json()["data"]["id"]

    monkeypatch.setattr(
        chat_service,
        "answer",
        lambda payload: ChatResponseDTO(
            answer=f"ok: {payload.query}",
            citations=[],
            model_used=payload.model_name or "unknown",
        ),
    )

    response = client.post(
        "/api/v1/chat/completions",
        json={
            "session_id": session_id,
            "query": "ping",
            "model_name": "gemma4:12b",
            "knowledge_base_ids": [],
        },
    )
    assert response.status_code == 200
    assert response.json()["data"]["answer"] == "ok: ping"


def test_docx_xlsx_ingest_and_delete() -> None:
    kb_response = client.post(
        "/api/v1/knowledge-bases",
        json={"name": "Parser KB", "description": "parser"},
    )
    assert kb_response.status_code == 200
    kb_id = kb_response.json()["data"]["id"]

    docx_response = client.post(
        f"/api/v1/knowledge-bases/{kb_id}/files/upload?enable_ocr=true",
        files={
            "file": (
                "plan.docx",
                io.BytesIO(_create_docx_bytes()),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert docx_response.status_code == 200
    docx_id = docx_response.json()["data"]["document"]["id"]

    xlsx_response = client.post(
        f"/api/v1/knowledge-bases/{kb_id}/files/upload?enable_ocr=true",
        files={
            "file": (
                "plan.xlsx",
                io.BytesIO(_create_xlsx_bytes()),
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        },
    )
    assert xlsx_response.status_code == 200

    chunks_response = client.get(f"/api/v1/knowledge-bases/{kb_id}/chunks")
    assert chunks_response.status_code == 200
    chunk_page = chunks_response.json()["data"]
    assert "items" in chunk_page
    chunk_texts = [item["content"] for item in chunk_page["items"]]
    assert any(
        "Edge deployment with local rag memory." in text for text in chunk_texts)
    assert any("Local retrieval" in text for text in chunk_texts)

    delete_response = client.delete(
        f"/api/v1/knowledge-bases/{kb_id}/files/{docx_id}")
    assert delete_response.status_code == 200
    files_response = client.get(f"/api/v1/knowledge-bases/{kb_id}/files")
    assert files_response.status_code == 200
    remaining_ids = [item["id"] for item in files_response.json()["data"]]
    assert docx_id not in remaining_ids


def test_upload_rejects_unsupported_extension() -> None:
    kb_response = client.post(
        "/api/v1/knowledge-bases",
        json={"name": "Reject Ext KB", "description": "validation"},
    )
    assert kb_response.status_code == 200
    kb_id = kb_response.json()["data"]["id"]

    upload_response = client.post(
        f"/api/v1/knowledge-bases/{kb_id}/files/upload?enable_ocr=true",
        files={
            "file": (
                "unsafe.exe",
                io.BytesIO(b"binary"),
                "application/octet-stream",
            )
        },
    )

    assert upload_response.status_code == 400
    assert "不支持该文件类型" in upload_response.json()["detail"]


def test_upload_rejects_oversized_file() -> None:
    kb_response = client.post(
        "/api/v1/knowledge-bases",
        json={"name": "Reject Size KB", "description": "validation"},
    )
    assert kb_response.status_code == 200
    kb_id = kb_response.json()["data"]["id"]

    upload_response = client.post(
        f"/api/v1/knowledge-bases/{kb_id}/files/upload?enable_ocr=true",
        files={
            "file": (
                "oversize.pdf",
                io.BytesIO(b"x" * (26 * 1024 * 1024)),
                "application/pdf",
            )
        },
    )

    assert upload_response.status_code == 400
    assert "文件超过大小限制" in upload_response.json()["detail"]


def test_upload_rejects_non_multimodal_model() -> None:
    kb_response = client.post(
        "/api/v1/knowledge-bases",
        json={"name": "Reject Model KB", "description": "validation"},
    )
    assert kb_response.status_code == 200
    kb_id = kb_response.json()["data"]["id"]

    upload_response = client.post(
        f"/api/v1/knowledge-bases/{kb_id}/files/upload?enable_ocr=true&model_name=gemma4:4b",
        files={
            "file": (
                "evidence.png",
                io.BytesIO(b"image-bytes"),
                "image/png",
            )
        },
    )

    assert upload_response.status_code == 400
    assert "当前模型不支持文件上传" in upload_response.json()["detail"]


def test_chat_attachment_upload_succeeds_for_multimodal_model(monkeypatch) -> None:
    monkeypatch.setattr(
        chat_service.ollama_service,
        "list_models",
        lambda: [{"name": "llava:7b", "size": 1,
                  "modified_at": "now", "digest": "vision"}],
    )
    monkeypatch.setattr(
        chat_attachment_service.parser_service,
        "extract_text",
        lambda *args, **kwargs: ("chat attachment body", "done"),
    )

    session_response = client.post(
        "/api/v1/sessions",
        json={"title": "Attachment Upload Session"},
    )
    assert session_response.status_code == 200
    session_id = session_response.json()["data"]["id"]

    upload_response = client.post(
        f"/api/v1/sessions/{session_id}/attachments?enable_ocr=true&model_name=llava:7b",
        files={
            "file": (
                "chat-note.txt",
                io.BytesIO(b"chat attachment body"),
                "text/plain",
            )
        },
    )

    assert upload_response.status_code == 200
    payload = upload_response.json()["data"]
    assert payload["session_id"] == session_id
    assert payload["file_name"] == "chat-note.txt"
    assert payload["attachment_type"] == "document"
    assert payload["status"] == "uploaded"
    assert "chat attachment body" in payload["extracted_text_preview"]


def test_chat_attachment_upload_rejects_non_multimodal_model() -> None:
    session_response = client.post(
        "/api/v1/sessions",
        json={"title": "Attachment Reject Session"},
    )
    assert session_response.status_code == 200
    session_id = session_response.json()["data"]["id"]

    upload_response = client.post(
        f"/api/v1/sessions/{session_id}/attachments?enable_ocr=true&model_name=gemma4:4b",
        files={
            "file": (
                "evidence.png",
                io.BytesIO(_tiny_png_bytes()),
                "image/png",
            )
        },
    )

    assert upload_response.status_code == 400
    assert "当前模型不支持文件上传" in upload_response.json()["detail"]


def test_chat_completions_binds_attachments_and_includes_message_payload(monkeypatch) -> None:
    monkeypatch.setattr(
        chat_service.ollama_service,
        "list_models",
        lambda: [{"name": "llava:7b", "size": 1,
                  "modified_at": "now", "digest": "vision"}],
    )
    monkeypatch.setattr(
        chat_attachment_service.parser_service,
        "extract_text",
        lambda file_name, *
        _args, **_kwargs: (f"parsed from {file_name}", "done"),
    )

    captured_chat_call: dict[str, object] = {}

    def _mock_chat(**kwargs: object) -> str:
        captured_chat_call.update(kwargs)
        return "mocked multimodal reply"

    monkeypatch.setattr(chat_service.ollama_service, "chat", _mock_chat)

    session_response = client.post(
        "/api/v1/sessions",
        json={"title": "Attachment Bind Session"},
    )
    assert session_response.status_code == 200
    session_id = session_response.json()["data"]["id"]

    text_upload = client.post(
        f"/api/v1/sessions/{session_id}/attachments?enable_ocr=true&model_name=llava:7b",
        files={
            "file": (
                "summary.txt",
                io.BytesIO(b"summary body"),
                "text/plain",
            )
        },
    )
    assert text_upload.status_code == 200
    text_attachment_id = text_upload.json()["data"]["id"]

    image_upload = client.post(
        f"/api/v1/sessions/{session_id}/attachments?enable_ocr=true&model_name=llava:7b",
        files={
            "file": (
                "diagram.png",
                io.BytesIO(_tiny_png_bytes()),
                "image/png",
            )
        },
    )
    assert image_upload.status_code == 200
    image_attachment_id = image_upload.json()["data"]["id"]

    chat_response = client.post(
        "/api/v1/chat/completions",
        json={
            "session_id": session_id,
            "query": "Please summarize the attachments",
            "model_name": "llava:7b",
            "knowledge_base_ids": [],
            "attachment_ids": [text_attachment_id, image_attachment_id],
        },
    )
    assert chat_response.status_code == 200
    assert chat_response.json()["data"]["answer"] == "mocked multimodal reply"

    chat_messages = captured_chat_call["messages"]
    assert isinstance(chat_messages, list)
    user_message = chat_messages[-1]
    assert isinstance(user_message, dict)
    assert "Attachment Context:" in str(user_message["content"])
    assert "parsed from summary.txt" in str(user_message["content"])
    assert "parsed from diagram.png" in str(user_message["content"])
    assert isinstance(user_message.get("images"), list)
    assert len(user_message["images"]) == 1

    session_detail = client.get(f"/api/v1/sessions/{session_id}")
    assert session_detail.status_code == 200
    messages = session_detail.json()["data"]["messages"]
    user_messages = [item for item in messages if item["role"] == "user"]
    assert len(user_messages) == 1
    attachments = user_messages[0]["attachments"]
    assert [item["file_name"]
            for item in attachments] == ["summary.txt", "diagram.png"]
    assert all(item["status"] == "linked" for item in attachments)


def test_chat_attachment_download_and_delete_flow(monkeypatch) -> None:
    monkeypatch.setattr(
        chat_service.ollama_service,
        "list_models",
        lambda: [{"name": "llava:7b", "size": 1,
                  "modified_at": "now", "digest": "vision"}],
    )
    monkeypatch.setattr(
        chat_attachment_service.parser_service,
        "extract_text",
        lambda *args, **kwargs: ("download me", "done"),
    )

    session_response = client.post(
        "/api/v1/sessions",
        json={"title": "Attachment Download Session"},
    )
    assert session_response.status_code == 200
    session_id = session_response.json()["data"]["id"]

    upload_response = client.post(
        f"/api/v1/sessions/{session_id}/attachments?enable_ocr=true&model_name=llava:7b",
        files={
            "file": (
                "download.txt",
                io.BytesIO(b"download me"),
                "text/plain",
            )
        },
    )
    assert upload_response.status_code == 200
    attachment_id = upload_response.json()["data"]["id"]

    download_response = client.get(
        f"/api/v1/chat/attachments/{attachment_id}/download",
    )
    assert download_response.status_code == 200
    assert download_response.content == b"download me"
    assert 'filename="download.txt"' in download_response.headers.get(
        "content-disposition", "")

    delete_response = client.delete(
        f"/api/v1/chat/attachments/{attachment_id}")
    assert delete_response.status_code == 200
    assert delete_response.json()["data"]["deleted"] is True

    missing_response = client.get(
        f"/api/v1/chat/attachments/{attachment_id}/download",
    )
    assert missing_response.status_code == 404


def test_knowledge_base_reindex_stats_and_chunk_pagination() -> None:
    kb_response = client.post(
        "/api/v1/knowledge-bases",
        json={"name": "Reindex KB", "description": "reindex"},
    )
    assert kb_response.status_code == 200
    kb_id = kb_response.json()["data"]["id"]

    upload_response = client.post(
        f"/api/v1/knowledge-bases/{kb_id}/files/upload?enable_ocr=true",
        files={
            "file": (
                "reindex.txt",
                io.BytesIO(
                    b"alpha beta gamma delta epsilon zeta eta theta iota"),
                "text/plain",
            )
        },
    )
    assert upload_response.status_code == 200
    document_id = upload_response.json()["data"]["document"]["id"]

    reindex_response = client.post(f"/api/v1/knowledge-bases/{kb_id}/reindex")
    assert reindex_response.status_code == 200
    task_id = reindex_response.json()["data"]["id"]

    for _ in range(40):
        task_response = client.get(f"/api/v1/tasks/{task_id}")
        assert task_response.status_code == 200
        task_payload = task_response.json()["data"]
        if task_payload["status"] == "done":
            break
    else:
        raise AssertionError("reindex task did not finish in time")

    stats_response = client.get(f"/api/v1/knowledge-bases/{kb_id}/stats")
    assert stats_response.status_code == 200
    stats = stats_response.json()["data"]
    assert stats["file_count"] == 1
    assert stats["chunk_count"] >= 1
    assert stats["token_count"] >= 1

    chunks_response = client.get(
        f"/api/v1/knowledge-bases/{kb_id}/chunks",
        params={"document_id": document_id, "offset": 0, "limit": 1},
    )
    assert chunks_response.status_code == 200
    payload = chunks_response.json()["data"]
    assert payload["limit"] == 1
    assert payload["offset"] == 0
    assert payload["total"] >= 1
    assert len(payload["items"]) == 1
    assert payload["items"][0]["document_id"] == document_id


def test_screenshot_source_question_uses_grounded_answer(monkeypatch) -> None:
    monkeypatch.setattr(
        chat_service.ollama_service,
        "list_models",
        lambda: [{"name": "qwen3.5:4b", "size": 1,
                  "modified_at": "now", "digest": "abc"}],
    )

    llm_call_count = {"value": 0}

    def _unexpected_llm_call(**_: object) -> str:
        llm_call_count["value"] += 1
        return "incorrect llm answer"

    monkeypatch.setattr(chat_service.ollama_service,
                        "chat", _unexpected_llm_call)
    monkeypatch.setattr(
        ingest_service.parser_service,
        "extract_text",
        lambda *args, **kwargs: (_screenshot_like_ocr_text(), "done"),
    )

    kb_response = client.post(
        "/api/v1/knowledge-bases",
        json={"name": "Screenshot KB", "description": "ocr"},
    )
    assert kb_response.status_code == 200
    kb_id = kb_response.json()["data"]["id"]

    upload_response = client.post(
        f"/api/v1/knowledge-bases/{kb_id}/files/upload?enable_ocr=true",
        files={
            "file": (
                "gpt-origin.png",
                io.BytesIO(b"fake-image-content"),
                "image/png",
            )
        },
    )
    assert upload_response.status_code == 200

    session_response = client.post(
        "/api/v1/sessions",
        json={"title": "Screenshot Session"},
    )
    assert session_response.status_code == 200
    session_id = session_response.json()["data"]["id"]

    chat_response = client.post(
        "/api/v1/chat/completions",
        json={
            "session_id": session_id,
            "query": "GPT 是生成预训练变换器的缩写，最初是在哪个论文中提出的",
            "model_name": "qwen3.5:4b",
            "knowledge_base_ids": [kb_id],
        },
    )
    assert chat_response.status_code == 200
    payload = chat_response.json()["data"]
    assert llm_call_count["value"] == 0
    assert "《通过生成预训练改善语言理解》" in payload["answer"]
    assert "2018" in payload["answer"]
    assert "Radford" in payload["answer"]
    assert "language_understanding_paper.pdf" in payload["answer"]
    assert payload["citations"]
    assert payload["citations"][0]["heading_path"] == "1.6 深入剖析GPT架构"
    assert payload["citations"][0]["matched_terms"]


def test_ollama_timeout_logs_are_queryable(monkeypatch) -> None:
    for path in config.logs_root.glob("app.log*"):
        path.write_text("", encoding="utf-8")

    monkeypatch.setattr(
        chat_service.ollama_service,
        "list_models",
        lambda: [{"name": "gemma4:12b", "size": 1,
                  "modified_at": "now", "digest": "abc"}],
    )

    def _raise_timeout(**_: object) -> str:
        raise httpx.ReadTimeout("timed out")

    monkeypatch.setattr(chat_service.ollama_service, "chat", _raise_timeout)

    trace_id = "trace-timeout-test"
    session_response = client.post(
        "/api/v1/sessions", json={"title": "Timeout Session"})
    session_id = session_response.json()["data"]["id"]

    response = client.post(
        "/api/v1/chat/completions",
        json={
            "session_id": session_id,
            "query": "trigger timeout",
            "model_name": "gemma4:12b",
            "knowledge_base_ids": [],
        },
        headers={"X-Trace-Id": trace_id},
    )
    assert response.status_code == 502

    logs_response = client.get(
        "/api/v1/logs",
        params={"module": "chat", "trace_id": trace_id, "limit": 20},
    )
    assert logs_response.status_code == 200
    events = logs_response.json()["data"]
    assert any(item["event"] == "chat.unexpected_error" for item in events)
    timeout_entries = [
        item for item in events if item["event"] == "chat.unexpected_error"]
    assert timeout_entries
    assert timeout_entries[-1]["trace_id"] == trace_id
    assert timeout_entries[-1]["context"]["session_id"] == session_id
    assert "timed out" in timeout_entries[-1]["message"].lower(
    ) or "timed out" in timeout_entries[-1].get("exception", "").lower()
